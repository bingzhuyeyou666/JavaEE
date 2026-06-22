from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


ROOT = Path(r"C:\Users\17661\Desktop\陌路寻阡\JavaEE")
WORK = ROOT / ".docx-work"
OUT = ROOT / "docs" / "参赛材料初稿"
OUT.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size=10.5, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(paragraph, text, size=10.5, bold=False, align=None):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25


def set_cell(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def unique_cells(row):
    result = []
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            result.append(cell)
    return result


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def insert_table_after(doc, paragraph, rows, cols, widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width
    return table


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p.getparent().remove(p._p)
    new_p.getparent().replace(new_p, p._p)
    if style:
        p.style = style
    set_paragraph_text(p, text)
    return p


def build_summary():
    doc = Document(WORK / "summary-template.docx")
    table = doc.tables[0]

    r = unique_cells(table.rows[0])
    set_cell(r[1], "【待填写】", 10.5)
    set_cell(r[3], "陌路寻阡智慧文旅综合服务平台", 10.5, True)

    r = unique_cells(table.rows[1])
    set_cell(r[1], "软件应用与开发", 10)
    set_cell(r[3], "Web应用／智慧文旅【以报名系统选项为准】", 10)

    set_cell(
        unique_cells(table.rows[2])[0],
        "作品简介（100字以内）：\n"
        "“陌路寻阡”是一套面向游客的智慧文旅平台，提供景点检索与详情、路线规划、景点预约热线查询、天气与地图、旅行广场、图生游记、足迹徽章及景点申报等功能，并配套后台审核和轮播内容管理。平台不提供在线预约、预约签到或核销服务",
        9,
    )
    set_cell(
        unique_cells(table.rows[3])[0],
        "创新描述（100字以内）：\n"
        "平台将景点浏览、路线排序、预约热线查询、旅行内容分享和后台审核串联为完整闭环，结合距离计算、标签推荐、天气降级服务与可选图生游记能力，在外部接口不可用时仍可保留核心业务运行",
        9,
    )
    set_cell(
        unique_cells(table.rows[4])[0],
        "特别说明\n"
        "1. 地图来源：配置密钥时使用百度地图开放平台，未配置时当前版本使用 OpenStreetMap 后备展示。正式提交前必须按最终参赛版本确认唯一地图来源并补充有效地图审图号；如无法满足赛事要求，应更换为符合要求的地图服务或素材\n"
        "2. 前期基础：作品在既有 Java Web、数据库与前端框架知识基础上独立设计开发，当前提交内容包括业务代码、界面、数据模型、测试与文档，具体起止时间由作者补充\n"
        "3. AI/AIGC 使用：开发过程中使用 OpenAI Codex 辅助部分代码审查、错误定位、注释整理、测试和文档初稿，输出均经作者人工复核修改；具体占比请作者按实际工作量填写【待填写：约____%】。作品中的图生游记功能可选调用阿里云百炼通义千问，正式部署前需确认账号、服务条款和内容合规情况",
        8.2,
    )

    author_header = unique_cells(table.rows[6])
    for index, name in enumerate(("【作者1】", "【作者2】", "【作者3】"), start=1):
        set_cell(author_header[index], name, 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    for row_index in range(7, 15):
        cells = unique_cells(table.rows[row_index])
        for cell in cells[1:4]:
            set_cell(cell, "____%", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(unique_cells(table.rows[14])[0], "说明：请按实际贡献填写，各作者各项目比例及总比例应符合赛事要求", 8.5)

    set_cell(
        unique_cells(table.rows[15])[1],
        "□作品创意  □理论指导  □技术方案  □实验场地  □硬件资源  □作品制作  □修改完善  □经费支持\n"
        "【请指导教师按实际情况勾选并补充】",
        9,
    )
    set_cell(unique_cells(table.rows[16])[1], "■Windows  □Linux  □macOS  □其他：浏览器及 Java／Node.js 开发环境", 9)
    set_cell(unique_cells(table.rows[17])[1], "■Windows  ■Linux  □macOS  □iOS  □Android  ■其他：现代浏览器", 9)
    set_cell(
        unique_cells(table.rows[18])[1],
        "IntelliJ IDEA、JDK 8、Maven、Spring Boot 2.7.18、Node.js、npm、Vite、React 19、H2 Database、Git、Chrome／Edge",
        9,
    )
    set_cell(
        unique_cells(table.rows[19])[1],
        "1. Spring Boot 2.7.18 Reference Documentation\n"
        "2. React Documentation\n"
        "3. 百度地图开放平台／Open-Meteo／OpenStreetMap 官方文档",
        9,
    )
    set_cell(
        unique_cells(table.rows[20])[1],
        "■素材压缩包  ■报告文档  ■演示视频  ■PPT  ■源代码  ■作品／部署文件\n"
        "说明：当前为整理计划，最终以实际上传文件为准",
        9,
    )

    files = [
        ("作品信息摘要", "按赛事模板填写的作品概况、分工、平台和文件清单", "自制"),
        ("设计和开发文档", "需求、总体设计、详细设计、测试及使用说明", "自制"),
        ("源代码压缩包", "后端、前端、配置及必要资源文件，不含 target 和 node_modules", "自制为主／依赖开源"),
        ("可运行程序", "Spring Boot 可运行 JAR、启动脚本及部署说明", "自制为主／依赖开源"),
        ("用户使用说明书", "面向评审的安装、登录、操作和常见问题说明", "自制"),
        ("代表性素材包", "轮播图、景点图片及界面素材，图片版权需逐项核验", "部分版权待核验"),
        ("演示PPT及PDF", "作品背景、功能、技术方案、创新点和演示流程", "待制作"),
        ("演示视频及答辩视频", "不超过赛事规定时长和大小的 MP4 文件", "待制作"),
    ]
    for row_index, (name, desc, copyright_text) in enumerate(files, start=23):
        cells = unique_cells(table.rows[row_index])
        set_cell(cells[1], f"文件名：{name}\n描述：{desc}", 8.2)
        set_cell(cells[2], "□已上传到网盘\n■未上传，待最终整理", 8.2)
        set_cell(cells[3], f"版权状态：{copyright_text}", 8.2)

    output = OUT / "02-作品信息摘要-陌路寻阡-初稿.docx"
    doc.save(output)
    return output


def add_section_table(doc, anchor, headers, data, widths=None):
    table = insert_table_after(doc, anchor, len(data) + 1, len(headers), widths)
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, 9, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], "D9E2F3")
    for row_idx, values in enumerate(data, start=1):
        for col_idx, value in enumerate(values):
            set_cell(table.rows[row_idx].cells[col_idx], value, 8.5)
    return table


def build_design():
    doc = Document(WORK / "design-template.docx")
    paragraphs = doc.paragraphs

    metadata = {
        5: "作品编号：【待填写】",
        6: "作品名称：陌路寻阡智慧文旅综合服务平台",
        7: "作者：【待填写】",
        8: "版本：V1.0",
        9: "日期：2026年6月【待填写】日",
    }
    for idx, text in metadata.items():
        set_paragraph_text(paragraphs[idx], text, 12, idx == 6, WD_ALIGN_PARAGRAPH.CENTER)

    requirements = (
        "本作品面向需要自主规划行程的游客和负责景点内容维护的管理人员。游客端需要解决景点信息分散、路线顺序难安排、"
        "预约信息查询不便、旅行记录不易沉淀等问题；管理端需要完成景点申报审核、首页轮播图维护和内容管理。"
        "系统采用浏览器访问方式，核心功能包括首页推荐、景点导览、路线规划、预约热线查询、旅行广场、图生游记、个人中心、"
        "景点申报与后台审核。非功能需求包括中文错误提示、外部接口降级、上传文件校验、会话鉴权、响应式界面和可重复部署。"
        "目标用户为普通游客、注册用户和管理员。与单一景点展示网站相比，本作品强调从浏览、规划、到访到分享的业务闭环；"
        "与重度商业旅游平台相比，系统规模适中、部署简单，适合校园项目展示和中小型文旅场景二次开发"
    )
    set_paragraph_text(paragraphs[16], requirements, 10.5)
    req_table = add_section_table(
        doc,
        paragraphs[16],
        ["角色", "主要需求", "对应模块"],
        [
            ["游客", "浏览景点、查看天气和地图、规划路线", "首页、景点导览、路线规划"],
            ["注册用户", "打卡、发布游记、收藏足迹", "足迹打卡、旅行广场、个人中心"],
            ["管理员", "审核申报、维护轮播和景点内容", "管理后台"],
        ],
        [Cm(2.2), Cm(7.5), Cm(5.5)],
    )

    overview = (
        "系统采用前后端分层的 Web 架构。浏览器端使用 React 19 和 Vite 构建交互界面，通过 HTTP 接口访问 Spring Boot 后端；"
        "后端按控制器、业务服务、数据访问和实体模型分层，使用 H2 数据库存储业务数据，并将上传图片保存到受控目录。"
        "天气、地图、语音和图像生成属于可选外部能力，均设置未配置提示或后备方案，避免外部服务故障导致核心功能不可用。"
        "用户接口与管理接口通过会话拦截器进行权限区分，统一异常处理负责把参数错误、数据错误和外部服务错误转换为中文提示"
    )
    set_paragraph_text(paragraphs[18], overview, 10.5)
    add_section_table(
        doc,
        paragraphs[18],
        ["层次", "组成", "职责"],
        [
            ["表现层", "React、Vite、CSS、浏览器", "页面展示、表单交互、状态管理和接口调用"],
            ["接口层", "Spring MVC Controller、拦截器", "接收请求、参数校验、登录与管理员权限控制"],
            ["业务层", "Service、推荐与路线计算", "景点、预约热线、打卡、广场、审核和外部服务编排"],
            ["数据层", "Spring Data JPA、H2、上传目录", "实体持久化、查询及图片文件管理"],
            ["外部服务", "百度地图、Open-Meteo、OpenStreetMap、可选AI接口", "地图、天气及扩展内容生成"],
        ],
        [Cm(2.2), Cm(6.0), Cm(7.0)],
    )

    details = (
        "1. 界面设计：顶部导航连接首页、景点导览、路线规划、旅行广场、图生游记、个人中心和景点申报。首页使用可维护轮播图展示核心入口；"
        "景点详情页展示图片、评分、天气、设施和地图；路线页支持选择2—5个景点并按交通方式生成结果；管理端通过文件上传维护轮播图片\n"
        "2. 数据设计：核心实体包括景点、设施、用户资料、签到记录、评价、广场帖子、评论、景点申报、轮播图和友善积分记录；预约热线直接使用景点电话字段展示。"
        "实体之间使用景点编号、用户标识和帖子编号关联，删除和审核操作在业务层校验状态\n"
        "3. 关键算法：路线智能排序采用近邻策略，使用经纬度距离寻找下一访问点；距离计算采用球面距离公式。签到时计算用户位置与景点坐标距离，"
        "超过500米时拒绝签到。推荐功能依据景点类型、标签、省份和用户行为进行加权。拥挤指数结合预约人数与景点容量估算\n"
        "4. 稳定性设计：天气服务优先读取配置接口，失败时切换 Open-Meteo；地图未配置时显示明确中文提示或后备地图。"
        "统一异常处理将类型转换、参数缺失、JSON格式和业务异常转换为中文信息，同时保留服务端日志供定位"
    )
    set_paragraph_text(paragraphs[20], details, 10)
    add_section_table(
        doc,
        paragraphs[20],
        ["数据对象", "主要内容", "用途"],
        [
            ["ScenicSpot", "名称、位置、坐标、分类、评分、容量、图片", "景点展示和路线计算"],
            ["ScenicSpot.phone／CheckInRecord", "景点热线、用户、景点、签到位置", "预约热线与到访记录"],
            ["SquarePost／SquareComment", "正文、图片、作者、点赞和评论", "旅行内容互动"],
            ["SpotSubmission", "申报人、景点资料、审核状态", "用户申报及管理员审核"],
            ["HeroSlide", "图片路径、标题、排序、启用状态", "首页轮播维护"],
        ],
        [Cm(4.0), Cm(6.0), Cm(5.2)],
    )

    tests = (
        "测试环境为 Windows、JDK 8、Maven、Node.js 和现代浏览器。后端执行 Maven 自动化测试，前端执行生产构建，"
        "并在本机 8080 端口进行页面与接口联调。当前自动化用例覆盖应用主页、景点接口、密钥不泄露、语音服务未配置状态、"
        "管理员轮播图上传、非法路线参数中文提示和非法 JSON 中文提示，共7项测试通过；前端生产构建成功。"
        "初步结果表明核心页面和接口可运行，错误信息能够以中文返回。正式提交前还需补充多浏览器、连续操作、较大图片上传、"
        "数据库备份恢复以及断网情况下的人工测试，并保存测试截图"
    )
    set_paragraph_text(paragraphs[22], tests, 10.5)
    add_section_table(
        doc,
        paragraphs[22],
        ["测试项", "预期结果", "初测结果"],
        [
            ["主页及景点接口", "页面可打开，景点数据可读取", "通过"],
            ["路线非法参数", "不输出英文堆栈，返回中文提示", "通过"],
            ["非法 JSON 请求", "返回可理解的中文格式错误", "通过"],
            ["轮播图片上传", "管理员可上传并生成可访问图片", "通过"],
            ["前端生产构建", "无编译错误并生成静态资源", "通过"],
            ["外部服务未配置", "提示配置状态，核心功能仍可用", "通过"],
        ],
        [Cm(4.2), Cm(7.3), Cm(3.5)],
    )

    usage = (
        "一、安装部署\n"
        "1. 安装 JDK 8 或兼容版本，确认 java 和 Maven 可用\n"
        "2. 在项目根目录执行 mvn clean package，生成可运行 JAR；前端资源会随构建流程打包\n"
        "3. 按需在环境变量中配置百度地图、天气、语音或阿里云百炼接口密钥，未配置时系统使用提示或后备能力\n"
        "4. 执行 java -jar target/travel-cloud-map-0.0.1-SNAPSHOT.jar，浏览器访问 http://localhost:8080\n"
        "二、基本使用\n"
        "1. 游客可浏览首页和景点详情并查询预约热线，注册登录后使用打卡、发布和个人中心功能\n"
        "2. 路线规划页选择2—5个景点，设置出发位置和交通方式后点击智能排序\n"
        "3. 管理员登录后台后可审核景点申报、维护景点，并通过上传图片方式切换首页轮播图\n"
        "三、注意事项\n"
        "正式提交材料不应包含 target、node_modules、IDE缓存和临时日志；图片素材需核验版权；地图来源与审图号必须按最终版本填写"
    )
    set_paragraph_text(paragraphs[24], usage, 10)

    summary = (
        "本项目完成了智慧文旅场景下从景点发现、路线安排、预约热线查询到内容分享和后台治理的基本闭环。"
        "开发过程中重点解决了前后端数据一致性、外部服务不可用时的降级、上传文件安全、会话权限和中文异常提示等问题。"
        "目前系统可在本地完成构建和运行，主要自动化测试通过。后续可继续完善地图合规配置、移动端适配、数据库迁移、"
        "更细致的推荐模型、压力测试和部署监控。参赛提交前应以实际运行版本重新截取界面图，并逐项核对文档描述、操作步骤和程序效果"
    )
    set_paragraph_text(paragraphs[26], summary, 10.5)

    references = (
        "[1] Spring Boot 2.7.18 Reference Documentation\n"
        "[2] React Documentation\n"
        "[3] Spring Data JPA Reference Documentation\n"
        "[4] 百度地图开放平台开发文档\n"
        "[5] Open-Meteo API Documentation\n"
        "[6] OpenStreetMap Documentation\n"
        "说明：正式提交时请补充实际查阅日期、网址及所使用第三方依赖的许可证信息"
    )
    set_paragraph_text(paragraphs[28], references, 10.5)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)

    output = OUT / "03-设计和开发文档-陌路寻阡-初稿.docx"
    doc.save(output)
    return output


def build_readme():
    text = """陌路寻阡参赛材料初稿说明

本目录中的文件根据赛事模板和当前项目代码初步整理，原始模板未被覆盖

一、需要作者补充的信息
1. 作品编号、学校、作者姓名、指导教师、联系电话等身份信息
2. 作者分工比例和指导教师实际作用
3. 项目真实开发起止时间
4. AI/AIGC 工具使用内容和实际比例
5. 最终地图来源及有效地图审图号
6. 所有图片、字体和第三方素材的版权或授权证明

二、提交前必须复核
1. 文档功能描述、操作步骤和最终程序保持一致
2. 使用最终程序重新截取界面图，避免旧截图
3. 执行后端测试和前端生产构建
4. 删除 target、node_modules、.idea、临时日志和缓存文件
5. 源代码包应能重新构建出相同的运行程序
6. 演示视频和答辩视频按赛事要求控制为 MP4、规定时长和大小

三、版权与合规提醒
轮播图和景点图片目前存在外部来源素材，不能在未核验版权的情况下直接填写“自制”
地图审图号和 AI 使用比例不能凭空填写，应以最终版本和真实开发过程为准
"""
    path = OUT / "README-提交前补充清单.txt"
    path.write_text(text, encoding="utf-8")
    return path


if __name__ == "__main__":
    outputs = [build_summary(), build_design(), build_readme()]
    for path in outputs:
        print(path)
